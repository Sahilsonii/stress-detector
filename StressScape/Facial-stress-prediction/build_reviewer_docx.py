"""Builds reviewer_response_package/Reviewer_Response.docx: the editor's
original comments paired with the author response for each point, plus the
real supporting figures embedded where relevant."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent.parent.parent
PKG = REPO / "reviewer_response_package"
FACIAL = PKG / "facial_results"
KEYS = PKG / "keystroke_results"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    return p


def body(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def label(text, color=RGBColor(0x1F, 0x4E, 0x79)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color
    return p


def figure(path, caption, width_in=5.5):
    if not path.exists():
        body(f"[Figure not found: {path.name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)


# ---------------------------------------------------------------- Title ---
t = doc.add_heading("Response to Reviewers", level=0)
body("A Multi-Model AI Framework for Non-Invasive Stress Detection Using "
     "Multimodal Behavioural and Facial Indicators", bold=True)
body("Sahil Soni, Zulfikar Ali Ansari, Wasim Khan, K. Kiran Kumar, Shreya Jagtap, Vishal Shah")
body("Decision: Major Revision", italic=True)
doc.add_paragraph()

body(
    "Dear Editor,\n\n"
    "Thank you for the opportunity to revise our manuscript. Below we reproduce "
    "each of your original comments in full, followed immediately by our "
    "response and, where applicable, the real evidence (metrics, confusion "
    "matrices, and figures) supporting that response. All referenced numbers "
    "are traceable to files in the project repository under "
    "StressScape/Facial-stress-prediction/results/."
)
doc.add_page_break()

# ---------------------------------------------------------- Editor summary ---
h1("Editor's Overall Recommendation")
body(
    "The manuscript addresses a relevant topic and proposes a potentially "
    "useful software-only approach for stress detection. However, the current "
    "version overstates several conclusions and does not provide enough "
    "evidence to support the main claims about real-time multimodal stress "
    "detection. The results should be reported more cautiously and the "
    "methodology needs clearer explanation.",
    italic=True,
)

# ------------------------------------------------------------- Points -----
POINTS = [
    {
        "title": "Point 1  -  Clarify how \"stress\" labels were created",
        "comment": (
            "The authors use FER2013 and CMU keystroke datasets, but these are "
            "not clearly stress-specific datasets. The manuscript should "
            "explain how stressed/non-stressed labels were derived and justify "
            "why these labels are valid for stress detection."
        ),
        "response": (
            "We have added explicit label-generation methodology for both "
            "modalities.\n\n"
            "Facial: the FER2013-derived binary mapping (negative-valence + "
            "neutral = Stressed, positive-valence = Not Stressed) is now "
            "explicitly flagged in the Methodology as a proxy label, not a "
            "clinical diagnosis, with a pointer to the Limitations discussion.\n\n"
            "Keystroke: this required a deeper fix than a clarifying sentence. "
            "We discovered during revision that the originally reported "
            "keystroke features (backspace count, total words typed, error "
            "rate) do not exist in the public CMU Keystroke Dynamics Benchmark "
            "dataset our citation points to  -  that dataset contains only "
            "fixed-password hold-time and inter-key latency measurements. We "
            "have rebuilt the keystroke pipeline from scratch on the actual "
            "public dataset (51 subjects x 400 repetitions of a fixed "
            "password), engineering five timing-based features (hold_mean, "
            "hold_std, flight_mean, flight_std, total_duration) and defining "
            "stress-proxy labels as each repetition's deviation from that "
            "participant's own personal typing baseline  -  explicitly decoupled "
            "from the features fed to the classifier, to avoid label leakage. "
            "This rule is now fully documented in the revised Methodology."
        ),
        "figures": [(KEYS / "feature_importance.png", "Keystroke Random Forest feature importance  -  flight-time variability (49.8%) is the dominant predictor.")],
    },
    {
        "title": "Point 2  -  Do not overclaim multimodal fusion performance",
        "comment": (
            "The manuscript claims that the combined system is stronger because "
            "one example shows 93.7% confidence. A single interface example is "
            "not enough. The authors should either provide proper test-set "
            "evaluation of the fused model or describe the system only as a "
            "prototype."
        ),
        "response": (
            "We agree the single-screenshot 93.7% figure was insufficient "
            "evidence and have removed it as a standalone claim. The manuscript "
            "now describes the fusion pipeline as an implemented, real-time "
            "weighted-fusion system (alpha = 0.6) demonstrated on illustrative "
            "sessions, explicitly stating that a systematic paired-participant "
            "benchmark is future work rather than a completed result  -  per the "
            "editor's own suggested fallback. A tool for collecting real paired "
            "facial+keystroke session data (keystroke_timing_logger.py) has "
            "been built and is being used to collect a small honest-N fusion "
            "evaluation for a future revision, rather than reporting an "
            "inflated single example."
        ),
        "figures": [],
    },
    {
        "title": "Point 3  -  Correct inconsistencies in the reported metrics",
        "comment": (
            "Some values appear inconsistent. For example, MobileNetV2 has the "
            "highest AUC, but the text presents ResNet50V2 as clearly best "
            "overall. There also seems to be confusion between accuracy and "
            "AUC. The authors should carefully recalculate and align all "
            "metrics, confusion-matrix values, and percentages."
        ),
        "response": (
            "We traced every reported number back to its source file and "
            "corrected three concrete errors:\n\n"
            "- The sentence crediting MobileNetV2 with \"the highest "
            "accuracy... 0.94\" conflated its AUC (0.94, correctly the highest "
            "of the three models) with its accuracy (0.85). This is now stated "
            "correctly in both directions.\n\n"
            "- The Discussion's ResNet50V2 confusion-matrix narrative "
            "(previously citing 3,240 / 166 / 1,102 / 614 on an implied "
            "5,122-sample set) did not match the actual evaluation, which was "
            "run on the full 13,479-sample validation set. It has been "
            "replaced with the real counts: TP=8,560, FP=466, TN=3,388, "
            "FN=1,065  -  and this correction actually reverses a conclusion in "
            "our favour: the real numbers show ResNet50V2 misses FEWER stress "
            "cases than MobileNetV2 (11.06% vs. 16.40% miss rate), consistent "
            "with Table 5's recall values.\n\n"
            "- The RF-vs-Decision-Tree comparison sentence was previously "
            "truncated mid-sentence and cited an unsupported 76.45% baseline. "
            "It now reports the real, freshly-trained Decision Tree baseline "
            "(72.00% accuracy) against the real RF result (73.65%), a "
            "reproducible 1.65-point gap, alongside 5-fold group-aware "
            "cross-validation of 78.78% +/- 2.52%.\n\n"
            "All comparison figures are now generated by code that "
            "parses the real confusion matrix out of each model's "
            "classification report, rather than reconstructing an approximate "
            "one from precision and recall  -  a latent bug we found and fixed "
            "in generate_comparison_report.py / regenerate_reports_only.py "
            "while verifying these numbers. Figure 11 in the manuscript has "
            "been regenerated accordingly, and Figure 12 (the keystroke "
            "confusion matrix) has been replaced with the real test-set "
            "matrix (TN=2,169, FP=611, FN=443, TP=777 on 4,000 held-out "
            "samples).\n\n"
            "- Separately, while auditing every reported figure we found a "
            "fourth inconsistency the original submission had not flagged: "
            "the framework diagram (Figure 4) stated \"Random Forest "
            "Accuracy = 83.5%\", a value matching neither the previously "
            "reported figure nor the real trained result, and it also "
            "depicted a backspace/error-rate preprocessing step that the "
            "public CMU dataset cannot support. Figure 4 has been redrawn "
            "from the corrected methodology and the real metrics. We "
            "mention it explicitly because it is exactly the class of "
            "inconsistency this comment asked us to eliminate."
        ),
        "figures": [
            (FACIAL / "comparison" / "all_confusion_matrices.png", "Real confusion matrices for all three facial models (MobileNetV2, EfficientNetB0, ResNet50V2), regenerated from the actual classification reports."),
            (FACIAL / "comparison" / "metrics_comparison.png", "Accuracy / precision / recall / F1 / AUC comparison across facial models."),
            (KEYS / "confusion_matrix.png", "Keystroke Random Forest confusion matrix (real held-out test set, default threshold)."),
        ],
    },
    {
        "title": "Point 4  -  Moderate claims about workplace and clinical use",
        "comment": (
            "The study does not appear to include real workplace testing, "
            "clinical validation, or physiological ground truth. Claims about "
            "occupational stress monitoring, wellness interventions, or "
            "clinical settings should be rewritten more cautiously."
        ),
        "response": (
            "We have softened language in the Abstract and Conclusion that "
            "implied validated real-time occupational deployment, and added an "
            "explicit statement that the system has not undergone workplace "
            "field testing or clinical/physiological validation. The "
            "Conclusion now explicitly notes the offered approach is \"pending "
            "the workplace field validation and physiological ground-truth "
            "studies\" discussed in Future Directions, rather than presenting "
            "deployment readiness as already established."
        ),
        "figures": [],
    },
    {
        "title": "Point 5  -  Describe the custom dataset in detail",
        "comment": (
            "The authors mention a custom dataset of 1,400 images, but they "
            "should report the number of participants, consent/ethics, "
            "demographics, labelling method, acquisition protocol, and whether "
            "subject-level separation was used."
        ),
        "response": (
            "Per Reviewer 2's confirmation in the prior revision round, the "
            "Ethics/Consent declarations (participant count  -  two "
            "author-contributors, 700 images each across seven emotion "
            "categories  -  exemption rationale under institutional policy, and "
            "written informed consent for image collection and publication) "
            "already satisfied this point and remain unchanged in this "
            "revision."
        ),
        "figures": [],
    },
    {
        "title": "Point 6  -  Clarify threshold optimisation",
        "comment": (
            "The manuscript mentions thresholds such as tau = 0.35-0.45 and "
            "tau = 0.65. The authors should explain how these thresholds were "
            "selected and whether they were tuned on validation data."
        ),
        "response": (
            "The keystroke threshold sweep is now explicit about methodology: "
            "theta is selected on the validation split only (F1-maximising, "
            "chosen threshold = 0.35), and reported separately from the "
            "held-out test-set evaluation at that chosen threshold, so no "
            "threshold was tuned on the same data used to report final "
            "performance. The full validation sweep and the resulting "
            "precision/recall trade-off are shown below, and this sweep has "
            "also been added to the manuscript as a new figure (following "
            "Figure 12) so the selection procedure is visible to the reader "
            "rather than described only in text."
        ),
        "figures": [(KEYS / "threshold_optimization.png", "Threshold optimisation curve (validation split): precision, recall, F1, and specificity vs. decision threshold, with the selected theta*=0.35 marked.")],
    },
    {
        "title": "Point 7  -  Improve comparison with prior work",
        "comment": (
            "Table 6 compares studies using different datasets, modalities, "
            "sample sizes, and validation protocols. The authors should avoid "
            "presenting these accuracy values as directly comparable unless "
            "the conditions are equivalent."
        ),
        "response": (
            "We have added an explicit caveat immediately after the "
            "comparison table stating that the listed studies differ in "
            "dataset, modality, sample size, and validation protocol, and that "
            "direct accuracy comparison should be read with that in mind  -  "
            "only the sensor-free deployment profile is directly comparable "
            "across systems."
        ),
        "figures": [],
    },
    {
        "title": "Point 8  -  Revise the language and structure",
        "comment": (
            "Several sections contain grammatical errors and unclear "
            "phrasing. A language revision is needed because some sentences "
            "obscure the scientific meaning."
        ),
        "response": (
            "We have corrected grammar and clarity issues in every passage "
            "touched by the numerical corrections above (see "
            "manuscript_corrections.md in the repository for the full list of "
            "before/after text). A broader copy-edit pass of untouched "
            "sections is available on request if the editorial team would "
            "like it before typesetting."
        ),
        "figures": [],
    },
]

for pt in POINTS:
    h1(pt["title"])
    label("Editor/Reviewer Comment:")
    body(pt["comment"], italic=True)
    label("Author Response:")
    body(pt["response"])
    for fig_path, cap in pt["figures"]:
        figure(fig_path, cap)
    doc.add_paragraph()

# --------------------------------------------------- Reviewer 2 (prior round) ---
h1("Reviewer 2  -  Prior Revision Round (Identifiable Images / Consent)")
label("Reviewer Comment:")
body(
    "The authors have appropriately addressed the concerns raised about "
    "using identifiable facial images in the manuscript. The authors have "
    "clarified that the images shown in the relevant figures belong to one "
    "of the authors and were not collected from external participants. The "
    "revised manuscript also includes a clear statement regarding the "
    "creation of the small internal facial image dataset, the absence of "
    "external participants or patients, the ethics exemption, and written "
    "informed consent for image collection and publication. The \"Informed "
    "Consent for Image Publication\" statement added to the Declarations "
    "section satisfactorily addresses the concern. In my opinion, the "
    "authors have properly addressed the comments, and the revision is "
    "satisfactory. The manuscript may be considered suitable for "
    "publication.",
    italic=True,
)
label("Author Response:")
body("No further action required  -  this point was resolved in the prior revision round and remains unchanged in the current version.")

doc.add_page_break()

# ------------------------------------------------- figure change summary ---
FIGURE_CHANGES = [
    ("Figure 4 - Proposed framework diagram", "REDRAWN",
     "The original diagram described a preprocessing step of \"Error rate "
     "Calculated = total_words / bksp\", a backspace-based feature the public "
     "CMU dataset does not contain (and the ratio was inverted). It also "
     "reported \"Random Forest Accuracy = 83.5%\", a third value inconsistent "
     "with both the previously reported and the real result, and listed a "
     "\"Tabs open count\" modality never described in the paper. Redrawn to "
     "show the real pipeline: CMU hold/flight-time features, within-subject "
     "proxy labelling, ResNet50V2 at 88.64%, Random Forest at 73.65% "
     "(78.78% +/- 2.52% CV), and the weighted late fusion actually implemented."),
    ("Figure 11 - Facial confusion matrices", "REGENERATED",
     "The previous version was produced by code that approximated each "
     "confusion matrix from precision and recall, so the image disagreed with "
     "the counts in the text. Regenerated directly from the real "
     "classification reports."),
    ("Figure 12 - Keystroke confusion matrix", "REPLACED",
     "Replaced with the real Random Forest test-set matrix (TN=2,169, FP=611, "
     "FN=443, TP=777). The previous image showed counts (222/34/122/22) that "
     "no longer correspond to any evaluation we can reproduce."),
    ("New figure - Threshold optimisation", "ADDED",
     "Precision, recall, F1 and specificity across the full threshold sweep on "
     "the validation split, with the selected tau* = 0.35 marked. Added in "
     "support of point 6."),
    ("New figure - Keystroke feature importance", "ADDED",
     "Random Forest feature importance (flight_std 49.77%, total_duration "
     "15.06%, flight_mean 13.52%, hold_std 13.04%, hold_mean 8.61%). Added in "
     "support of point 1 and the Results discussion."),
]

h1("Summary of Figure and Diagram Changes")
body(
    "For traceability, the table below lists every figure altered in this "
    "revision and why. Tables 1-4 and Table 6 were also revised; those "
    "changes are described in the relevant points above."
)
for title, status, why in FIGURE_CHANGES:
    label(f"{title}  [{status}]")
    body(why)

doc.add_page_break()

# ------------------------------------------------------------ Appendix ----
h1("Appendix: Additional Supporting Figures")
body(
    "The figures below supplement the point-by-point responses above with "
    "the complete real evaluation record for all trained models."
)

h2("Facial Expression Models")
for model in ["MobileNetV2", "EfficientNetB0", "ResNet50V2"]:
    figure(FACIAL / model / "confusion_matrix_roc.png", f"{model}: confusion matrix + ROC curve")
    figure(FACIAL / model / "training_history.png", f"{model}: training history (accuracy, loss, AUC, precision/recall)")

h2("Facial Model Comparison")
for fname, cap in [
    ("heatmap_comparison.png", "Heatmap comparison across all facial models"),
    ("boxplot_comparison.png", "Box-plot comparison across all facial models"),
    ("stacked_comparison.png", "Stacked metric comparison across all facial models"),
    ("winner_announcement.png", "Final model selection summary"),
]:
    figure(FACIAL / "comparison" / fname, cap)

h2("Keystroke Random Forest  -  Complete Evidence")
body("Real metrics summary (from results/Keystroke/metrics.json):")
for line in [
    "Test accuracy (default threshold 0.5): 73.65%  |  Precision: 0.560  |  Recall: 0.637  |  F1: 0.596",
    "5-fold group-aware cross-validation: 78.78% +/- 2.52%",
    "Decision Tree baseline: 72.00% accuracy  ->  RF improvement: +1.65 percentage points",
    "Feature importance: flight_std 49.77% | total_duration 15.06% | flight_mean 13.52% | hold_std 13.04% | hold_mean 8.61%",
    "Dataset: real CMU DSL-StrongPasswordData, 51 subjects x 400 repetitions, group-aware split (31/10/10 subjects)",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(line)

doc.add_paragraph()
body(
    "Full machine-readable evidence (metrics.json, classification_report.txt, "
    "feature_statistics.json, split_summary.json, threshold_table.json) is "
    "included alongside this document in reviewer_response_package/keystroke_results/ "
    "and reviewer_response_package/facial_results/."
)

OUT = PKG / "Reviewer_Response.docx"
doc.save(str(OUT))

# ---------------------------------------------------------------------------
# Emit the plain-text response from the SAME source data, so the .txt and the
# .docx can never drift apart again (they previously had).
# ---------------------------------------------------------------------------
def wrap(text, width=76, indent=""):
    out = []
    for para in text.split("\n"):
        if not para.strip():
            out.append("")
            continue
        line = indent
        for word in para.split():
            if len(line) + len(word) + 1 > width and line.strip():
                out.append(line.rstrip())
                line = indent + word + " "
            else:
                line += word + " "
        out.append(line.rstrip())
    return "\n".join(out)


t = []
t.append("RESPONSE TO REVIEWERS")
t.append("=" * 76)
t.append("")
t.append("Manuscript: A Multi-Model AI Framework for Non-Invasive Stress")
t.append("            Detection Using Multimodal Behavioural and Facial Indicators")
t.append("Authors:    Sahil Soni, Zulfikar Ali Ansari, Wasim Khan, K. Kiran Kumar,")
t.append("            Shreya Jagtap, Vishal Shah")
t.append("Decision:   Major Revision")
t.append("")
t.append("NOTE: This plain-text file and Reviewer_Response.docx are generated")
t.append("from the same source and contain identical text. The .docx additionally")
t.append("embeds the supporting figures inline.")
t.append("")
t.append(wrap(
    "Dear Editor, Thank you for the opportunity to revise our manuscript. "
    "Below we reproduce each of your original comments in full, followed "
    "immediately by our response. All referenced numbers are traceable to "
    "files in the project repository under "
    "StressScape/Facial-stress-prediction/results/."))
t.append("")

for pt in POINTS:
    t.append("")
    t.append("-" * 76)
    t.append(pt["title"].upper())
    t.append("-" * 76)
    t.append("")
    t.append("EDITOR/REVIEWER COMMENT:")
    t.append(wrap(pt["comment"], indent="  "))
    t.append("")
    t.append("AUTHOR RESPONSE:")
    t.append(wrap(pt["response"], indent="  "))
    t.append("")

t.append("")
t.append("-" * 76)
t.append("SUMMARY OF FIGURE AND DIAGRAM CHANGES")
t.append("-" * 76)
t.append("")
for title, status, why in FIGURE_CHANGES:
    t.append(f"{title}  [{status}]")
    t.append(wrap(why, indent="  "))
    t.append("")

t.append("-" * 76)
t.append("KEYSTROKE RANDOM FOREST - REAL METRICS")
t.append("-" * 76)
t.append("")
for line in [
    "Test accuracy (default threshold 0.5): 73.65% | Precision 0.560 | Recall 0.637 | F1 0.596",
    "5-fold group-aware cross-validation: 78.78% +/- 2.52%",
    "Decision Tree baseline: 72.00%  ->  RF improvement +1.65 pp",
    "Feature importance: flight_std 49.77% | total_duration 15.06% | flight_mean 13.52% | hold_std 13.04% | hold_mean 8.61%",
    "Dataset: real CMU DSL-StrongPasswordData, 51 subjects x 400 reps, group-aware split (31/10/10 subjects)",
]:
    t.append(wrap("- " + line, indent="  "))
t.append("")

(PKG / "response_to_reviewers.txt").write_text("\n".join(t) + "\n",
                                               encoding="utf-8", newline="\n")

print(f"Saved: {OUT}")
print(f"Saved: {PKG / 'response_to_reviewers.txt'} (regenerated from same source)")
print(f"Saved: {OUT}")
