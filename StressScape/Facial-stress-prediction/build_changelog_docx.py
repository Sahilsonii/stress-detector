"""Builds reviewer_response_package/Manuscript_Change_Log.docx

A section-by-section before/after record of every change between the ORIGINAL
unedited manuscript and the revised one, produced by actually diffing the two
.tex files (not written from memory), plus the figure/diagram changes with the
old and new images shown side by side.
"""

import re
import json
import difflib
import pathlib

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = pathlib.Path(__file__).resolve().parent.parent.parent
ORIG_DIR = pathlib.Path('C:/Users/Sahil/Downloads/Stressscape___sahil_soni')
NEW_DIR = REPO / 'overleaf_submission'
PKG = REPO / 'reviewer_response_package'

OLD_TEX = ORIG_DIR / 'main.tex'
NEW_TEX = NEW_DIR / 'main.tex'

RED = 'FDECEA'
GREEN = 'EAF7EC'
GREY = 'F4F4F4'

SEC_RE = re.compile(r'\\(?:sub)?section\*?\{([^}]*)\}')

# --------------------------------------------------------------- diff -------
def blocks(t):
    out, sec = [], 'Preamble / Abstract'
    for para in re.split(r'\n\s*\n', t):
        p = para.strip()
        if not p:
            continue
        m = SEC_RE.search(p)
        if m:
            sec = m.group(1)
        out.append((sec, p))
    return out


old_raw = OLD_TEX.read_text(encoding='utf-8', errors='replace').replace('\r\n', '\n')
new_raw = NEW_TEX.read_text(encoding='utf-8', errors='replace').replace('\r\n', '\n')
ob, nb = blocks(old_raw), blocks(new_raw)
sm = difflib.SequenceMatcher(None, [b[1] for b in ob], [b[1] for b in nb], autojunk=False)

hunks = []
for tag, i1, i2, j1, j2 in sm.get_opcodes():
    if tag == 'equal':
        continue
    sec = ob[i1][0] if i1 < len(ob) else (nb[j1][0] if j1 < len(nb) else '')
    hunks.append({'section': sec,
                  'old': [ob[i][1] for i in range(i1, i2)],
                  'new': [nb[j][1] for j in range(j1, j2)]})

# ------------------------------------------------- curated annotations ------
# keyed by hunk index -> (where in paper, what changed, which editor point(s))
NOTES = {
 0: ("Abstract",
     "Dataset name corrected from \"JAFFE dataset\" to \"FER2013 public benchmark "
     "dataset\" (every other section already said FER2013). The keystroke sentence "
     "was rewritten: the feature description changed from error rate / backspace "
     "frequency to hold-time and flight-time typing rhythm, and the unsupported "
     "84.56% / 0.92 / 0.8567 figures were replaced with the real trained results. "
     "The closing sentence no longer asserts a 93.7% fused confidence.",
     "Points 1, 2, 3"),
 1: ("Introduction - contributions list",
     "The third contribution bullet previously claimed \"achieving 0.92 precision\". "
     "Rewritten to describe the real public-dataset pipeline, the documented "
     "within-subject proxy-labelling method, and group-aware cross-validation, "
     "without quoting an unsupported precision figure.",
     "Points 1, 3"),
 2: ("Methodology - Keystroke Dynamics Dataset",
     "Dataset description replaced. The original said \"51 participants, totalling "
     "400 keys\" with participant ID / backspace count / total words / error rate "
     "columns. The public CMU benchmark contains none of those columns; it is 51 "
     "subjects x 400 repetitions of a fixed password with per-key hold (H), "
     "down-down (DD) and up-down (UD) latencies. Now described accurately, with "
     "the 20,400 total repetitions stated.",
     "Point 1"),
 3: ("Methodology - Table 1, Feature Engineering, Stress Label Generation "
     "(Eq. 2-3), Table 2",
     "The largest single change. Table 1 replaced (backspace/words/error-rate rows "
     "-> real hold/flight/duration feature rows). Feature Engineering rewritten to "
     "the five timing features actually computed. The error-rate equation was "
     "replaced by the feature-definition equation. Stress Label Generation replaced: "
     "the old rule thresholded error rate at 0.25; the new rule is a within-subject "
     "relative typing-irregularity score thresholded at the 70th percentile of the "
     "TRAIN split only, with an explicit note that the classifier never sees the "
     "baseline or the score, to avoid label leakage. Table 2 replaced with real "
     "statistics over all 20,400 repetitions.",
     "Points 1, 6"),
 4: ("Methodology - Table 3 (group-aware split)",
     "Approximate counts (\"30-31 participants, ~1200 samples\") replaced with the "
     "actual split: 31 / 10 / 10 subjects and 12,400 / 4,000 / 4,000 samples.",
     "Point 1"),
 5: ("Methodology - Keystroke Model Architectures",
     "The sentence listing what Random Forest handles well changed from \"typing "
     "patterns, backspace counts, and error rates\" to \"hold-time and flight-time "
     "statistics\", matching the real feature set.",
     "Point 1"),
 6: ("Methodology - Feature Importance Ranking",
     "The ranking line changed from \"error_rate > backspace_count > total_words\" "
     "to the real measured ranking with percentages.",
     "Points 1, 3"),
 7: ("Methodology - Threshold Optimization Analysis + Table 4",
     "The text now states explicitly that the threshold was selected on the "
     "VALIDATION split only (F1-maximising, tau* = 0.35) and reports the held-out "
     "test performance separately, so no threshold was tuned on the data used for "
     "final reporting. Table 4 replaced with the real sweep and re-captioned "
     "\"(validation split)\".",
     "Point 6"),
 8: ("Results - Keystroke Dynamics Performance, Figure 12, two new figures, "
     "Real-Time Interface",
     "Cross-validation result corrected from 84.56% +/- 1.68% to the real 78.78% "
     "+/- 2.52% (group-aware), and the feature-importance sentence rewritten. "
     "Figure 12 re-pointed from the old screenshot to the real keystroke confusion "
     "matrix. TWO NEW FIGURES inserted here (threshold optimisation, feature "
     "importance). The Real-Time Interface paragraph was rewritten: the 93.7% "
     "fused-confidence claim and the \"stronger than single-method\" conclusion "
     "were removed, and the section now states the interface is illustrative and "
     "that a paired benchmark is future work.",
     "Points 2, 3, 6"),
 9: ("Discussion - Facial Expression Modality, Keystroke Dynamics Modality, "
     "Multimodal Fusion Validity, Comparison, Critical Limitations",
     "Four separate corrections. (a) The MobileNetV2 sentence conflated AUC 0.94 "
     "with accuracy - now stated correctly (AUC 0.94, accuracy 0.85 vs ResNet50V2 "
     "0.89). (b) The ResNet50V2 confusion-matrix narrative (3,240/166/1,102/614 on "
     "an implied 5,122 samples) replaced with the real counts on 13,479 samples; "
     "this REVERSES the original conclusion, since ResNet50V2 actually misses fewer "
     "stress cases than MobileNetV2. (c) The RF-vs-Decision-Tree sentence was "
     "truncated mid-sentence and cited an unsupported 76.45% baseline and \"60 "
     "trees\" - now the real 73.65% vs 72.00% gap with 100 trees. (d) Fusion "
     "Validity reframed as implemented-but-not-yet-benchmarked. (e) Table 6 gained "
     "a comparability caveat, and Critical Limitations gained a sentence on "
     "keystroke proxy-label validity.",
     "Points 2, 3, 4, 7"),
 10: ("Conclusion",
      "The keystroke claim changed from 84.56% to the real 73.65% test / 78.78% CV. "
      "The Random Forest sentence no longer claims \"0.92 precision and 0.8567 "
      "recall\". Added that the training pipeline and evaluation artefacts are in "
      "the repository, that fusion is implemented but not yet benchmarked, and that "
      "the workplace claim is pending field and physiological validation.",
      "Points 2, 3, 4"),
 11: ("Data Availability",
      "GitHub URL corrected - it pointed at github.com/Sahilsonii/stress-detector, "
      "which is not this project's repository. Added an explicit pointer to the "
      "keystroke model, training script and evaluation artefacts so every reported "
      "number is traceable.",
      "Point 3"),
}

# ---------------------------------------------------- figure change record ---
FIGURES = [
 ("Figure 4", "Methodology (framework diagram)", "REDRAWN",
  ORIG_DIR / 'all models compariosn' / 'unnamed.png',
  NEW_DIR / 'all models compariosn' / 'unnamed.png',
  "The original diagram contained three problems. (1) Its preprocessing box read "
  "\"Error rate Calculated = total_words/ bksp\" - a backspace-based feature the "
  "public CMU dataset does not contain, and the ratio was also inverted (error "
  "rate would be backspaces divided by words, not the reverse). (2) It reported "
  "\"Random Forest Accuracy = 83.5%\", a THIRD value matching neither the "
  "originally reported 84.56% nor the real trained 73.65%. (3) It listed a \"Tabs "
  "open count\" modality that appears nowhere in the paper. Redrawn to show the "
  "real pipeline, with ResNet50V2 at 88.64% and Random Forest at 73.65% "
  "(78.78% +/- 2.52% CV) and the weighted late fusion actually implemented in code."),
 ("Figure 11", "Results (facial confusion matrices)", "REGENERATED",
  ORIG_DIR / 'all models compariosn' / 'all_confusion_matrices.png',
  NEW_DIR / 'all models compariosn' / 'all_confusion_matrices.png',
  "The original was produced by report code that ESTIMATED each confusion matrix "
  "from precision and recall (tp = stressed x recall, tn = not_stressed x accuracy) "
  "rather than loading the real matrix, so the image disagreed with the counts in "
  "the text. That bug was fixed in generate_comparison_report.py and "
  "regenerate_reports_only.py, and the figure regenerated from the real "
  "classification reports."),
 ("Figure 12", "Results (keystroke confusion matrix)", "REPLACED",
  ORIG_DIR / 'all models compariosn' / 'WhatsApp Image 2025-11-13 at 11.31.37.png',
  NEW_DIR / 'all models compariosn' / 'keystroke_confusion_matrix.png',
  "The original image was a screenshot showing counts 222 / 34 / 122 / 22, which do "
  "not correspond to any reproducible evaluation. Replaced with the real Random "
  "Forest test-set matrix: TN=2,169, FP=611, FN=443, TP=777 on 4,000 held-out "
  "samples (subject-disjoint). The file was also renamed from \"WhatsApp Image "
  "2025-11-13 at 11.31.37.png\" to keystroke_confusion_matrix.png."),
 ("New figure (after Fig. 12)", "Results (threshold optimisation)", "ADDED",
  None,
  NEW_DIR / 'all models compariosn' / 'keystroke_threshold_optimization.png',
  "Did not exist before. Shows precision, recall, F1 and specificity across the "
  "full threshold sweep on the validation split, with the selected tau* = 0.35 "
  "marked. Added because editor point 6 asked how thresholds were selected and "
  "whether they were tuned on validation data - this makes the answer visible "
  "rather than text-only."),
 ("New figure (after the above)", "Results (keystroke feature importance)", "ADDED",
  None,
  NEW_DIR / 'all models compariosn' / 'keystroke_feature_importance.png',
  "Did not exist before. Shows the measured Random Forest feature importance "
  "(flight_std 49.77%, total_duration 15.06%, flight_mean 13.52%, hold_std 13.04%, "
  "hold_mean 8.61%). The percentages were previously quoted in the text with no "
  "supporting figure."),
]

UNCHANGED = [
 ("Table 5 (facial model comparison)",
  "Needed no correction - the values (MobileNetV2 0.85/0.94, EfficientNetB0 "
  "0.81/0.75, ResNet50V2 0.89/0.83) already matched the real metrics.json files. "
  "Only the surrounding prose was wrong, not the table."),
 ("Figures 1, 2, 3 (dataset preview and class distributions)",
  "Unchanged. Verified against dataset_count_report.json - the class counts and "
  "totals (28,300 original train / 40,096 balanced train / 53,571 total) are correct."),
 ("Figures 5-10 (per-model training history and confusion matrix + ROC)",
  "Unchanged. These were already generated from the real training runs."),
 ("Figures 13-15 (real-time interface screenshots)",
  "Images unchanged; only their captions were revised to drop the 93.7% claim. "
  "NOTE: if the 93.7% value is legible inside the screenshot pixels themselves, "
  "these images should be retaken so they do not contradict the corrected caption."),
 ("Ethics Approval, Informed Consent for Image Publication, Authors' "
  "Contributions, Funding, Competing Interests",
  "Unchanged. Reviewer 2 confirmed in the prior revision round that the consent "
  "and ethics statements were satisfactory."),
 ("Related Work, Model Architectures (MobileNetV2 / EfficientNetB0 / "
  "ResNet50V2 descriptions), Two-Phase Training Protocol, Algorithm 1",
  "Unchanged - these described the facial pipeline accurately and required no "
  "correction."),
]

# --------------------------------------------------------------- doc --------
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def codeblock(text, fill):
    """monospace shaded single-cell table - survives Word reflow cleanly"""
    text = re.sub(r'\n{3,}', '\n\n', text.strip())
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    shade(c, fill)
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8)
    for row in t.rows:
        for cc in row.cells:
            cc.width = Inches(6.5)
    return t


def para(text, italic=False, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic, r.bold = italic, bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def tag(text, color=RGBColor(0x1F, 0x4E, 0x79), size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def picture(path, caption, width=5.6):
    if path is None or not pathlib.Path(path).is_file():
        para(f'[image not available: {path}]', italic=True, size=9)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)


# ------------------------------------------------------------- title -------
doc.add_heading('Manuscript Change Log', level=0)
para('A Multi-Model AI Framework for Non-Invasive Stress Detection Using '
     'Multimodal Behavioural and Facial Indicators', bold=True)
para('Section-by-section record of every change between the original submitted '
     'manuscript and the revised version.')
doc.add_paragraph()

tag('How this document was produced')
para(f'The two LaTeX sources were compared programmatically (paragraph-level '
     f'diff), so this list is derived from the files themselves rather than '
     f'written from recollection.', size=10)
codeblock(f'ORIGINAL : {OLD_TEX}\nREVISED  : {NEW_TEX}\n\n'
          f'Paragraph blocks : {len(ob)} original -> {len(nb)} revised\n'
          f'Changed regions  : {len(hunks)} text/table hunks\n'
          f'Figure changes   : {len(FIGURES)} (tracked separately - a text diff '
          f'cannot see image content)', GREY)
doc.add_paragraph()

tag('Summary')
summary = doc.add_table(rows=1, cols=3)
summary.style = 'Light Grid Accent 1'
hdr = summary.rows[0].cells
for i, h in enumerate(['#', 'Section / location', 'Editor point(s)']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
for i, h in enumerate(hunks):
    where, _, pts = NOTES.get(i, (h['section'], '', '-'))
    row = summary.add_row().cells
    row[0].text = f'T{i + 1}'
    row[1].text = where
    row[2].text = pts
for fig in FIGURES:
    row = summary.add_row().cells
    row[0].text = fig[0].split()[0][:2] + '-fig'
    row[1].text = f'{fig[1]}  [{fig[2]}]'
    row[2].text = 'Point 3' if fig[2] != 'ADDED' else 'Points 1, 6'
doc.add_page_break()

# ------------------------------------------------- PART A: text changes ----
doc.add_heading('Part A  -  Text and Table Changes', level=1)
para('For each region: where it sits in the paper, what changed and why, then '
     'the original text followed by the revised text, verbatim from the LaTeX '
     'source.', italic=True, size=10)

for i, h in enumerate(hunks):
    where, what, pts = NOTES.get(i, (h['section'], 'Changed.', '-'))
    doc.add_heading(f'T{i + 1}.  {where}', level=2)
    tag(f'Editor point(s): {pts}', color=RGBColor(0xB2, 0x5F, 0x00), size=9.5)
    tag('What changed and why')
    para(what, size=10)

    tag('ORIGINAL (before revision)', color=RGBColor(0xA5, 0x2A, 0x2A))
    if h['old']:
        codeblock('\n\n'.join(h['old']), RED)
    else:
        para('(nothing here in the original - this content is new)',
             italic=True, size=9)

    tag('REVISED (current)', color=RGBColor(0x1E, 0x6B, 0x2E))
    if h['new']:
        codeblock('\n\n'.join(h['new']), GREEN)
    else:
        para('(removed in the revision)', italic=True, size=9)
    doc.add_paragraph()

doc.add_page_break()

# ---------------------------------------------- PART B: figure changes -----
doc.add_heading('Part B  -  Figure and Diagram Changes', level=1)
para('A text diff cannot see image content, so these are recorded separately. '
     'The original and revised images are shown together for each change.',
     italic=True, size=10)

for name, where, status, oldp, newp, why in FIGURES:
    doc.add_heading(f'{name}  -  {where}   [{status}]', level=2)
    tag('Why it changed')
    para(why, size=10)
    if oldp is not None:
        tag('ORIGINAL image', color=RGBColor(0xA5, 0x2A, 0x2A))
        picture(oldp, f'BEFORE  -  {pathlib.Path(oldp).name}')
    tag('REVISED image' if oldp is not None else 'NEW image',
        color=RGBColor(0x1E, 0x6B, 0x2E))
    picture(newp, f'AFTER  -  {pathlib.Path(newp).name}')
    doc.add_paragraph()

doc.add_page_break()

# ------------------------------------------- PART C: deliberately unchanged -
doc.add_heading('Part C  -  Checked and Deliberately Left Unchanged', level=1)
para('These were examined during the revision and required no change. Listed so '
     'the record is complete and it is clear they were not overlooked.',
     italic=True, size=10)
for what, why in UNCHANGED:
    tag(what)
    para(why, size=10)
    doc.add_paragraph()

OUT = PKG / 'Manuscript_Change_Log.docx'
doc.save(str(OUT))
print(f'Saved: {OUT}')
print(f'  {len(hunks)} text/table hunks, {len(FIGURES)} figure changes, '
      f'{len(UNCHANGED)} unchanged entries')
